
# ---- CELL 0 (code) ----
import pandas as pd
import numpy as np
import os
from collections.abc import Iterable
from typing import List
# --- mount (safe) ---
import urllib.request
from pathlib import Path

BASE_URL = "https://raw.githubusercontent.com/FelixZhan/AtyAN/main/"
HELPER_FILES = [
    "analysis_utils.py",
    "requirements.txt",
    "BP1234-ONSET.csv",
]

for filename in HELPER_FILES:
    dest = Path(filename)
    if dest.exists():
        print(f"{filename} already present, skipping download.")
        continue
    print(f"Downloading {filename}...")
    urllib.request.urlretrieve(f"{BASE_URL}{filename}", dest)

print("Helper files are ready.")
df = pd.read_csv("BP1234-ONSET.csv", low_memory=False)
ID_COL = "id"

def cols_exist(df: pd.DataFrame, cols: Iterable[str]) -> List[str]:
    return [c for c in cols if c in df.columns]
# Ensure ID exists
if ID_COL not in df.columns:
    for cand in ["id","ID"]:
        if cand in df.columns:
            ID_COL = cand; break


# ---- CELL 1 (code) ----
import pandas as pd
import numpy as np
import re
import statsmodels.api as sm
import statsmodels.formula.api as smf
from statsmodels.stats.multitest import fdrcorrection


# ---- CELL 2 (code) ----
def build_risk_prodromal(df: pd.DataFrame):
    """
    Create baseline risk-factor and prodromal variables in RAW units.
    Mirrors the Aim 1 ANOVA setup.
    """
    # ---- Risk factors (baseline) ----
    risk_candidates = [
        "w1tii",     # thin-ideal internalization
        "w1bs",      # body dissatisfaction
        "w1dres",    # dietary restraint
        "w1socf",    # social functioning
        "w1dep",     # depression
        "w1intbmi",  # BMI (kg/m^2)
        "w1age",     # age in years
    ]
    for col in risk_candidates:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors="coerce")
    RISK = [c for c in risk_candidates if c in df.columns]

    # ---- Prodromals (baseline, RAW) ----
    # Binge eating
    df["BE_w1"] = pd.to_numeric(df.get("w1ede1a"), errors="coerce")

    # Compensatory behaviors (max across 4 items)
    cb_cols = [c for c in ["w1ed8a", "w1ed9a", "w1ed10a", "w1ed11a"] if c in df.columns]
    if cb_cols:
        cb_mat = df[cb_cols].apply(pd.to_numeric, errors="coerce")
        df["CB_w1"] = cb_mat.max(axis=1, skipna=True)
        # if *all* four items missing → NaN
        df.loc[cb_mat.notna().sum(axis=1) == 0, "CB_w1"] = np.nan
    else:
        df["CB_w1"] = np.nan

    # Cognitive prodromal features
    df["WSO_w1"]  = pd.to_numeric(df.get("w1ed15a"), errors="coerce")  # wt/shape overvaluation
    df["FEAR_w1"] = pd.to_numeric(df.get("w1ed17a"), errors="coerce")  # fear of weight gain
    df["FAT_w1"]  = pd.to_numeric(df.get("w1ed19a"), errors="coerce")  # feeling fat

    # Lower-than-expected BMI (LEB)
    mbmi_pct = pd.to_numeric(df.get("w1mbmi_pct"), errors="coerce")
    df["LEB_w1"] = np.clip(90.0 - mbmi_pct, 0, None) / 90.0

    PRODROMAL = ["BE_w1", "CB_w1", "WSO_w1", "FEAR_w1", "FAT_w1", "LEB_w1"]
    OUTCOMES = [c for c in (RISK + PRODROMAL) if c in df.columns]

    return df, OUTCOMES


# ---- CELL 3 (code) ----
def has_cols_any(df: pd.DataFrame, prefixes):
    """
    Return a boolean mask: 'present' if any of the family columns (fan.00, pan.00, etc.)
    are non-zero numeric or a 'truthy' string.

    This mirrors the logic in the notebook:
      - look for PREFIX.00 / .01 / .02
      - numeric != 0 counts as present
      - strings TRUE/T/YES/Y/1 count as present
    """
    cols = []
    for p in prefixes:
        pat = re.compile(rf"^{re.escape(p)}\.(?:00|01|02)$", flags=re.I)
        cols.extend([c for c in df.columns if pat.match(c)])
    if not cols:
        return pd.Series(False, index=df.index)

    sub = df[cols].copy()
    for c in cols:
        if sub[c].dtype == object:
            s = sub[c].astype(str).str.strip().str.upper()
            truthy = s.isin(["TRUE", "T", "YES", "Y", "1"])
            falsy  = s.isin(["FALSE", "F", "NO", "N", "0", ""])
            sub[c] = np.where(truthy, 1,
                              np.where(falsy, 0, np.nan))
        else:
            sub[c] = pd.to_numeric(sub[c], errors="coerce")

    present = (sub != 0) & sub.notna()
    return present.any(axis=1)


# ---- CELL 4 (markdown) ----
# building group masks. defining obesity first based on BMI percentile (95+) and then adding that comparison to the ANOVAs

# ---- CELL 5 (code) ----
import pandas as pd
import numpy as np

# Download CDC BMI-for-age z-score reference
CDC_BMI_URL = "https://www.cdc.gov/growthcharts/data/zscore/bmiagerev.csv"
cdc_bmi = pd.read_csv(CDC_BMI_URL)

def compute_obesity_mask(df: pd.DataFrame, cdc_bmi: pd.DataFrame) -> pd.Series:
    """
    Return a boolean Series (index = df.index) indicating obesity:
    BMI-for-age >= 95th percentile, using CDC BMI-for-age reference.

    Assumptions:
      - df has w1age (YEARS) and w1intbmi (kg/m^2).
      - Use CDC rows with Sex == 2 (female).
      - For each participant, Agemos = w1age * 12.
      - Use the closest explicit Agemos row by rounding DOWN.
      - Missing age or BMI → not obese (False).
    """

    # Make sure required columns are present
    required_cols = ["Sex", "Agemos", "P95"]
    missing = [c for c in required_cols if c not in cdc_bmi.columns]
    if missing:
        raise ValueError(f"CDC BMI file missing required columns: {missing}")

    # Filter to female rows, being robust to typing/whitespace
    fem = cdc_bmi.copy()
    fem["Sex_str"] = fem["Sex"].astype(str).str.strip()
    fem = fem[fem["Sex_str"] == "2"]

    if fem.empty:
        print("Warning: no female (Sex==2) rows found in CDC BMI table; "
              "returning all participants as non-obese.")
        return pd.Series(False, index=df.index)

    fem = fem.sort_values("Agemos")
    ages_grid = pd.to_numeric(fem["Agemos"], errors="coerce").to_numpy()
    p95_grid  = pd.to_numeric(fem["P95"],    errors="coerce").to_numpy()

    # Participant age (years → months) and BMI
    age_years  = pd.to_numeric(df.get("w1age"),    errors="coerce")
    bmi_values = pd.to_numeric(df.get("w1intbmi"), errors="coerce")
    age_months = age_years * 12.0

    # Valid rows: have both age and BMI
    valid = age_months.notna() & bmi_values.notna()

    # Preallocate BMI-95th-percentile series
    bmi_p95 = pd.Series(np.nan, index=df.index)

    if valid.any():
        valid_ages = age_months[valid].to_numpy()

        # For each valid age, find index of greatest ages_grid <= age_months (rounding DOWN)
        idx = np.searchsorted(ages_grid, valid_ages, side="right") - 1

        # Clamp indices to [0, len(ages_grid)-1]
        idx[idx < 0] = 0
        idx[idx >= len(ages_grid)] = len(ages_grid) - 1

        bmi_p95_valid = p95_grid[idx]
        bmi_p95.loc[valid] = bmi_p95_valid

    # Obesity = BMI >= age-specific 95th percentile, with non-missing age & BMI
    mask_obese = valid & (bmi_values >= bmi_p95)
    mask_obese = mask_obese.fillna(False)

    return mask_obese


# ---- CELL 6 (code) ----
def build_group_masks(df: pd.DataFrame):
    # ---- Family "any" masks (AN/BN/BED/PU) ----
    mask_AN_any  = has_cols_any(df, ["fan", "pan"])
    mask_BN_any  = has_cols_any(df, ["fbn", "pbn"])
    mask_BED_any = has_cols_any(df, ["fbe", "pbe"])
    mask_PU_any  = has_cols_any(df, ["fpu", "ppu"])

    # ---- AtyAN / AAN from baseline onset flag ----
    if "w1ONSET-FULL" in df.columns:
        mask_atyAN = df["w1ONSET-FULL"].astype(str).str.strip().str.upper().eq("TRUE")
    else:
        mask_atyAN = pd.Series(False, index=df.index)

    # ---- Full-threshold only, excluding AAN ----
    mask_AN_full  = has_cols_any(df, ["fan"]) & ~mask_atyAN
    mask_BN_full  = has_cols_any(df, ["fbn"]) & ~mask_atyAN
    mask_BED_full = has_cols_any(df, ["fbe"]) & ~mask_atyAN
    mask_PU_full  = has_cols_any(df, ["fpu"]) & ~mask_atyAN

    # ---- Hierarchical exclusivity: AN > BN > BED > PU ----
    # (one ED per person, as in the notebook)
    mask_AN_only   = mask_AN_full
    mask_BN_only   = mask_BN_full  & ~mask_AN_only
    mask_BED_only  = mask_BED_full & ~mask_AN_only & ~mask_BN_only
    mask_PU_only   = mask_PU_full  & ~mask_AN_only & ~mask_BN_only & ~mask_BED_only

    # ---- AAN restricted to those without a full-threshold ED ----
    mask_AAN = (
        mask_atyAN &
        ~mask_AN_only &
        ~mask_BN_only &
        ~mask_BED_only &
        ~mask_PU_only
    )

    # ---- 10% weight-loss group ("WL10_noCog") ----
    # From your sample code: w1HEALTHY-WL & not AAN
    if "w1HEALTHY-WL" in df.columns:
        mask_WL10 = df["w1HEALTHY-WL"].fillna(False).astype(bool) & ~mask_AAN
    else:
        mask_WL10 = pd.Series(False, index=df.index)

    # ---- Obesity (BMI-for-age >= 95th percentile) ----
    mask_Obese = compute_obesity_mask(df, cdc_bmi)

    # ---- Healthy / no ED group ----
    # No AAN and no full AN/BN/BED/PU. (You can decide whether to also
    # exclude Obese from "Healthy" by uncommenting the last line.)
    mask_healthy = (
        ~mask_AAN &
        ~mask_AN_only &
        ~mask_BN_only &
        ~mask_BED_only &
        ~mask_PU_only
        # & ~mask_Obese    # <- uncomment if you want Healthy = non-obese, no-ED
    )

    groups = {
        "AAN":     mask_AAN,
        "Healthy": mask_healthy,
        "WL10":    mask_WL10,
        "AN":      mask_AN_any,
        "BN":      mask_BN_any,
        "BE":      mask_BED_any,
        "PU":      mask_PU_any,
        "Obese":   mask_Obese,   # BMI-for-age ≥ 95th percentile (female CDC chart)
    }
    return groups


# ---- CELL 7 (code) ----
def anova_one_contrast(df, outcomes, mask_g1, mask_g2, name_g1, name_g2):
    """
    Run 1-way ANOVAs for AAN vs a comparator group across multiple outcomes.
    Returns means, SDs, F, p, and eta^2 for each outcome.
    """
    mask_g1 = mask_g1.reindex(df.index).fillna(False).astype(bool)
    mask_g2 = mask_g2.reindex(df.index).fillna(False).astype(bool)
    mask_g2 = mask_g2 & ~mask_g1  # ensure no overlap

    group = pd.Series(np.nan, index=df.index, dtype=object)
    group.loc[mask_g1] = name_g1
    group.loc[mask_g2] = name_g2
    in_contrast = group.notna()

    rows = []
    for y in outcomes:
        if y not in df.columns:
            continue

        ysel = df.loc[in_contrast, [y]]
        if ysel.shape[1] != 1:
            continue

        yv = pd.to_numeric(ysel.iloc[:, 0], errors="coerce")
        gv = group.loc[in_contrast]

        ok = yv.notna()
        yv = yv[ok]
        gv = gv[ok]

        n1 = int((gv == name_g1).sum())
        n2 = int((gv == name_g2).sum())
        m1 = float(yv[gv == name_g1].mean()) if n1 else np.nan
        s1 = float(yv[gv == name_g1].std(ddof=1)) if n1 > 1 else np.nan
        m2 = float(yv[gv == name_g2].mean()) if n2 else np.nan
        s2 = float(yv[gv == name_g2].std(ddof=1)) if n2 > 1 else np.nan

        if n1 >= 2 and n2 >= 2:
            dfx = pd.DataFrame({y: yv.values, "group": gv.values})
            model = smf.ols(f"{y} ~ C(group)", data=dfx).fit()
            aov = sm.stats.anova_lm(model, typ=2)
            F = float(aov.loc["C(group)", "F"])
            p = float(aov.loc["C(group)", "PR(>F)"])
            ss_between = float(aov.loc["C(group)", "sum_sq"])
            ss_total = ss_between + float(aov.loc["Residual", "sum_sq"])
            eta2 = ss_between / ss_total if ss_total > 0 else np.nan
        else:
            F = p = eta2 = np.nan

        rows.append({
            "comparison": f"{name_g1}_vs_{name_g2}",
            "group1": name_g1,
            "group2": name_g2,
            "outcome": y,
            "n_group1": n1,
            "mean_group1": m1,
            "sd_group1": s1,
            "n_group2": n2,
            "mean_group2": m2,
            "sd_group2": s2,
            "F": F,
            "p": p,
            "eta2": eta2,
        })

    return pd.DataFrame(rows).sort_values(["comparison", "outcome"])


# ---- CELL 8 (code) ----


# ---- CELL 9 (code) ----
def bh_q(pvals: pd.Series) -> pd.Series:
    """
    Benjamini-Hochberg FDR q-values for a vector of p-values.
    """
    p = pd.to_numeric(pvals, errors="coerce").values
    mask = np.isfinite(p)
    q = np.full_like(p, np.nan, dtype=float)
    if mask.sum():
        _, qvals = fdrcorrection(p[mask], alpha=0.05, method="indep")
        q[mask] = qvals
    return pd.Series(q, index=pvals.index)


# ---- CELL 10 (code) ----
def run_aan_baseline_anovas(df: pd.DataFrame) -> pd.DataFrame:
    """
    Full one-way ANOVA setup for baseline AAN vs:
      - Healthy (no ED)
      - WL10 (10% weight-loss, w1HEALTHY-WL, no AAN)
      - AN, BN, BE, PU

    Returns a tidy DataFrame with:
      outcome, n/mean/SD by group, F, p, eta2, p_raw, q_all, q_by_outcome.
    """
    # Build baseline variables
    df, outcomes = build_risk_prodromal(df)

    # Group masks
    masks = build_group_masks(df)
    mask_AAN = masks["AAN"]

    comparisons = {
        "Healthy": masks["Healthy"],
        "WL10":    masks["WL10"],
        "Obese":   masks["Obese"],
        "AN": masks["AN"],
        "BN": masks["BN"],
        "BE": masks["BE"],
        "PU": masks["PU"],
    }

    all_rows = []
    for name, mask_comp in comparisons.items():
        res = anova_one_contrast(df, outcomes, mask_AAN, mask_comp, "AAN", name)
        all_rows.append(res)

    if not all_rows:
        return pd.DataFrame()

    result = pd.concat(all_rows, ignore_index=True)

    # ---- Multiple-testing corrections ----
    result["p_raw"] = pd.to_numeric(result["p"], errors="coerce")

    # FDR across ALL tests (all outcomes, all comparisons)
    result["q_all"] = bh_q(result["p_raw"])

    # Optional: FDR within each outcome (e.g., across comparators)
    result["q_by_outcome"] = (
        result.groupby("outcome", group_keys=False)["p_raw"].apply(bh_q)
    )

    return result


# ---- CELL 11 (code) ----
import re 
masks = build_group_masks(df)
aan_anova_tbl = run_aan_baseline_anovas(df)

# Peek at a few rows
print(aan_anova_tbl.head().to_string(index=False))

# ---- CELL 12 (markdown) ----
# time to print!

# ---- CELL 13 (code) ----
!pip install python-docx

import pandas as pd
import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---- CELL 14 (code) ----
# 1) Outcome label mapping (your requested display names)
OUTCOME_LABELS = {
    # Risk factors
    "w1bs":     "Body dissatisfaction",
    "w1dep":    "Negative affect",
    "w1dres":   "Dieting",
    "w1intbmi": "BMI",
    "w1socf":   "Psychosocial functioning",
    "w1tii":    "Thin-ideal internalization",

    # Prodromals (baseline)
    "BE_w1":    "Binge eating",
    "CB_w1":    "Compensatory behaviors",
    "FAT_w1":   "Feeling fat",
    "FEAR_w1":  "Fear of weight gain",
    "LEB_w1":   "Lower-than-expected BMI",
    "WSO_w1":   "Weight/shape overvaluation",

    # Just in case the suffix got dropped in the table:
    "FAT":      "Feeling fat",
    "FEAR":     "Fear of weight gain",
    "LEB":      "Lower-than-expected BMI",
    "WSO":      "Weight/shape overvaluation",
}

# 2) Desired outcome order (risk factors first, then prodromal symptoms)
OUTCOME_ORDER = [
    "w1bs",
    "w1dep",
    "w1dres",
    "w1intbmi",
    "w1socf",
    "w1tii",
    "BE_w1",
    "CB_w1",
    "FAT_w1",
    "FEAR_w1",
    "LEB_w1",
    "WSO_w1",
]

def format_pq(x):
    """Format p- and q-values: if < .001 and >0 → '<0.001'; else 3 decimals."""
    if pd.isna(x):
        return ""
    try:
        x = float(x)
    except (TypeError, ValueError):
        return ""
    if 0 < x < 0.001:
        return "<0.001"
    return f"{x:.3f}"

def fmt_num(x, nd=2):
    """Format numeric values with nd decimals; empty string if NaN."""
    if pd.isna(x):
        return ""
    try:
        return f"{float(x):.{nd}f}"
    except (TypeError, ValueError):
        return ""

def prepare_comparison_table(aan_anova_tbl, comparison_name):
    """
    Take the full ANOVA table and return a tidy DataFrame for one comparison:
    - Filters to a single comparison
    - Drops w1age
    - Applies outcome label mapping
    - Orders rows by OUTCOME_ORDER
    - Formats nothing yet (keeps raw numbers)
    """
    if comparison_name not in aan_anova_tbl["comparison"].unique():
        return None

    sub = aan_anova_tbl[aan_anova_tbl["comparison"] == comparison_name].copy()

    # Drop w1age entirely
    sub = sub[sub["outcome"] != "w1age"]

    # Restrict to outcomes we know / care about (optional, but keeps tidy)
    sub = sub[sub["outcome"].isin(OUTCOME_ORDER)]

    # Outcome sort index
    order_map = {name: i for i, name in enumerate(OUTCOME_ORDER)}
    sub["outcome_sort"] = sub["outcome"].map(order_map)
    sub = sub.sort_values("outcome_sort")

    # Add a human-readable label
    sub["Outcome"] = sub["outcome"].map(OUTCOME_LABELS).fillna(sub["outcome"])

    return sub


# ---- CELL 15 (code) ----
# Decide which FDR column to use as "q" (global FDR across tests is typical)
q_col = "q_all" if "q_all" in aan_anova_tbl.columns else "q_by_outcome"

# Decide which eta2 column to use (in case you renamed to partial eta2)
eta_col = "eta2"
for candidate in ["eta2_p", "eta2_partial", "partial_eta2"]:
    if candidate in aan_anova_tbl.columns:
        eta_col = candidate
        break

# Map comparison code → human-friendly table title
COMPARISON_TITLES = {
    "AAN_vs_Healthy": "Table 1. Baseline risk factors and prodromal symptoms: Atypical AN vs healthy participants",
    "AAN_vs_WL10":    "Table 2. Baseline risk factors and prodromal symptoms: Atypical AN vs 10% weight-loss without cognitive ED symptoms",
    "AAN_vs_AN": "Table 3. Baseline risk factors and prodromal symptoms: Atypical AN vs threshold/subthreshold AN",
    "AAN_vs_BN": "Table 4. Baseline risk factors and prodromal symptoms: Atypical AN vs threshold/subthreshold BN",
    "AAN_vs_BE": "Table 5. Baseline risk factors and prodromal symptoms: Atypical AN vs threshold/subthreshold BED",
    # PI calls this PD; data may use PU_full; the name here is just text in the document
    "AAN_vs_PU": "Table 6. Baseline risk factors and prodromal symptoms: Atypical AN vs purging disorder",
    # If you later add obese comparison, e.g. "AAN_vs_Obese", you can extend here:
    "AAN_vs_Obese":  "Table 7. Baseline risk factors and prodromal symptoms: Atypical AN vs obese participants",
}

# Order in which to create tables, filtered to only those actually present
comparison_order = [
    "AAN_vs_Healthy",
    "AAN_vs_WL10",
    "AAN_vs_AN",
    "AAN_vs_BN",
    "AAN_vs_BE",
    "AAN_vs_PU",
    "Obese"
]

available_comparisons = [c for c in comparison_order
                         if c in aan_anova_tbl["comparison"].unique()]

# Create Word document
doc = Document()

for comp in available_comparisons:
    sub = prepare_comparison_table(aan_anova_tbl, comp)
    if sub is None or sub.empty:
        continue

    # Figure out which groups are being compared (for column labels / header)
    # They should all be the same within a comparison
    g1 = sub["group1"].iloc[0]
    g2 = sub["group2"].iloc[0]

    # Title paragraph
    title_text = COMPARISON_TITLES.get(
        comp,
        f"Baseline risk factors and prodromal symptoms: {g1} vs {g2}"
    )
    p = doc.add_paragraph(title_text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Small spacer paragraph
    doc.add_paragraph("")

    # Create table:
    # Columns: Outcome, g1 n, g1 M, g1 SD, g2 n, g2 M, g2 SD, F, p, q, partial eta^2
    n_rows = sub.shape[0] + 1
    n_cols = 11
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Outcome"
    hdr_cells[1].text = f"{g1} n"
    hdr_cells[2].text = f"{g1} M"
    hdr_cells[3].text = f"{g1} SD"
    hdr_cells[4].text = f"{g2} n"
    hdr_cells[5].text = f"{g2} M"
    hdr_cells[6].text = f"{g2} SD"
    hdr_cells[7].text = "F"
    hdr_cells[8].text = "p"
    hdr_cells[9].text = "q (FDR)"
    hdr_cells[10].text = "partial η²"

    # Fill table rows
    for i, (_, row) in enumerate(sub.iterrows(), start=1):
        cells = table.rows[i].cells

        cells[0].text = str(row["Outcome"])

        cells[1].text = str(int(row["n_group1"])) if not pd.isna(row["n_group1"]) else ""
        cells[2].text = fmt_num(row["mean_group1"], 2)
        cells[3].text = fmt_num(row["sd_group1"], 2)

        cells[4].text = str(int(row["n_group2"])) if not pd.isna(row["n_group2"]) else ""
        cells[5].text = fmt_num(row["mean_group2"], 2)
        cells[6].text = fmt_num(row["sd_group2"], 2)

        cells[7].text = fmt_num(row["F"], 3)
        cells[8].text = format_pq(row.get("p_raw", row.get("p", np.nan)))
        cells[9].text = format_pq(row.get(q_col, np.nan))
        cells[10].text = fmt_num(row.get(eta_col, np.nan), 3)

    # Add a blank paragraph between tables
    doc.add_paragraph("")

# Save and download
docx_filename = "AAN_baseline_comparisons_tables.docx"
doc.save(docx_filename)

# In Colab, this will trigger a direct download to your laptop.
try:
    from google.colab import files
    files.download(docx_filename)
except ImportError:
    print(f"Saved {docx_filename} in the current working directory.")


# ---- CELL 16 (markdown) ----
# if in VS code in new extension
there's no way to download in the normal way right now

# ---- CELL 17 (code) ----
ls

# ---- CELL 18 (code) ----
import base64, zipfile, os

# ---- 1) Make sure these match the files you created ----
files_to_export = [
    # "anova_AAN_vs_baseline_groups_RAW_with_FDR.csv",
    "AAN_baseline_comparisons_tables.docx",
]

# Optional: check they exist
for fn in files_to_export:
    if not os.path.exists(fn):
        print("Missing:", fn)

zip_name = "aan_exports.zip"

# ---- 2) Create a zip with both files ----
with zipfile.ZipFile(zip_name, "w") as z:
    for fn in files_to_export:
        if os.path.exists(fn):
            z.write(fn)

# ---- 3) Base64-encode the zip and print as text ----
with open(zip_name, "rb") as f:
    b64 = base64.b64encode(f.read()).decode("ascii")

# To avoid giant single-line outputs, print in chunks
chunk_size = 8000
for i in range(0, len(b64), chunk_size):
    print(b64[i:i+chunk_size])


# ---- CELL 19 (code) ----


# ---- CELL 20 (code) ----

